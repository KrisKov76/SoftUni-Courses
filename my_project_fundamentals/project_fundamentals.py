import docx
# # Creating a Word Document with Python-docx
# document = docx.Document()
# document.save("d:\krischo.docx")
#
# # Adding Text to the Document
# document = docx.Document()
# paragraph = document.add_paragraph("Eto, napisah text!.")
# document.save("d:\krischo.docx")

## multiple paragraphs to the document by repeating the add_paragraph method
document = docx.Document()
paragraph = document.add_paragraph("Имало някога едно малко сладко момиченце. Всеки го обиквал от пръв поглед, но най-много го обичала баба му, която всеки път се чудела какво да даде на детето. Веднъж му подарила шапчица от червено кадифе, която му стояла тъй хубаво, че то не искало да носи друга и затова хората почнали да го наричат Червената шапчица..")
paragraph.add_run(" Един ден майка му рекла:")
document.add_paragraph("— Червена шапчице, ето ти малко козунак и едно шише вино. Занеси ги на баба си, защото тя е болна и немощна и трябва да се подкрепи. Тръгни, докато слънцето не е почнало да прежуря, върви мирно и тихо и не се отбивай от пътя, защото може да паднеш, да строшиш шишето и да оставиш баба си без вино. А щом влезеш в стаята й, недей забравя да поздравиш с „Добро утро“ и не любопитствувай да видиш какво става край тебе.")
document.save("d:\krischo.docx")